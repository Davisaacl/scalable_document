# app/pipeline/parsers.py
"""
sniff_mime: based on extensions

parse: generates the doc_id and the Blocks list
"""
import os, uuid, json, mimetypes
from typing import List
from .document_models import ParsedDocument, Block

# Parsing dependencies
import pdfplumber
from docx import Document as Docx

# Supported extensions
EXT_TO_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".json": "application/json",
}

def sniff_mime(path: str) -> str:
    # 1) For each extension (cross-platform)
    ext = os.path.splitext(path)[1].lower()
    if ext in EXT_TO_MIME:
        return EXT_TO_MIME[ext]
    # 2) Defining mimetypes
    mt, _ = mimetypes.guess_type(path)
    return mt or "application/octet-stream"

def parse(path: str) -> ParsedDocument:
    mime = sniff_mime(path)
    doc_id = str(uuid.uuid4())  
    if mime == EXT_TO_MIME[".pdf"]:
        blocks: List[Block] = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                blocks.append(Block(text=text, meta={"type": "page", "page": i + 1}))
        return ParsedDocument(doc_id, path, mime, title=None, blocks=blocks)

    if mime == EXT_TO_MIME[".docx"]:
        d = Docx(path)
        blocks: List[Block] = []
        for p in d.paragraphs:
            if p.text.strip():
                blocks.append(Block(text=p.text, meta={"type": "paragraph"}))
        return ParsedDocument(doc_id, path, mime, title=None, blocks=blocks)

    if mime == EXT_TO_MIME[".json"]:
        # Can use UTF-8 or without BOM
        with open(path, "r", encoding="utf-8-sig") as f:
            obj = json.load(f)

        lines: List[str] = []
        def walk(prefix, val):
            if isinstance(val, dict):
                for k, v in val.items():
                    walk(f"{prefix}.{k}" if prefix else k, v)
            elif isinstance(val, list):
                for j, v in enumerate(val):
                    walk(f"{prefix}[{j}]", v)
            else:
                lines.append(f"{prefix}: {val}")
        walk("", obj)
        blocks = [Block(text="\n".join(lines), meta={"type": "json"})]
        return ParsedDocument(doc_id, path, mime, title=None, blocks=blocks)

    # Not supported if we reach here
    raise ValueError(f"Unsupported type: {mime} (path={path})")
